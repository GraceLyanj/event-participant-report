import sys
import os
import re
import unicodedata
from collections import Counter
# Check required packages first and give a clear fix if missing
def _check_imports():
    missing = []
    try:
        import pandas
    except ImportError:
        missing.append("pandas")
    try:
        import matplotlib
    except ImportError:
        missing.append("matplotlib")
    try:
        import docx
    except ImportError:
        missing.append("python-docx")
    if missing:
        pip_names = {"pandas": "pandas", "matplotlib": "matplotlib", "python-docx": "python-docx", "docx": "python-docx"}
        to_install = [pip_names.get(m, m) for m in missing]
        to_install = list(dict.fromkeys(to_install))  # dedupe (docx -> python-docx)
        exe = sys.executable
        print("Missing required package(s):", ", ".join(missing))
        print("Using Python:", exe)
        print("Run this in the same terminal you use to run the script:")
        print(f'  "{exe}" -m pip install {" ".join(to_install)}')
        sys.exit(1)

_check_imports()

import pandas as pd
import io
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def get_proportions_df(series):
    """Return (counts, proportions, dataframe) for the series."""
    counts = series.value_counts(dropna=False)
    proportions = counts / counts.sum()
    df = pd.DataFrame({'Category': counts.index.astype(str), 'Count': counts.values, 'Proportion': proportions.values})
    return counts, proportions, df


def find_first_matching_column(df, candidates):
    """Return first existing column from candidates (case-insensitive exact, then contains)."""
    if df is None or df.empty:
        return None
    normalized = {c.strip().lower(): c for c in df.columns}
    for candidate in candidates:
        key = candidate.strip().lower()
        if key in normalized:
            return normalized[key]
    for candidate in candidates:
        key = candidate.strip().lower()
        for col in df.columns:
            low = col.strip().lower()
            if key in low:
                return col
    return None


def find_eid_column(df):
    """Participant primary key: EID column (Toolkit exports vary by header spelling)."""
    if df is None or df.empty:
        return None
    norm = {c.strip().lower(): c for c in df.columns}
    for key in ("eid", "ut eid", "student eid", "empl eid", "emp eid", "utexas eid"):
        if key in norm:
            return norm[key]
    col = find_first_matching_column(df, ["EID", "UT EID", "Student EID", "Empl EID"])
    if col:
        return col
    for c in df.columns:
        low = c.strip().lower()
        if low == "eid" or ("eid" in low and "pseudo" not in low):
            return c
    return None


def dedupe_participants_by_eid(df):
    """
    Ensure one row per EID so every chart/table uses the same person-level denominator.
    Rows with missing EID are kept distinct (each counts separately). Duplicate EIDs keep the first row.
    """
    eid_col = find_eid_column(df)
    if not eid_col:
        print("No EID-like column found; rows are not deduplicated by person.")
        return df.reset_index(drop=True)
    out = df.copy()

    def key_for(val, row_idx):
        if val is None or (isinstance(val, float) and pd.isna(val)):
            return f"__missing_eid_{row_idx}"
        t = str(val).strip()
        if t == "" or t.lower() == "nan":
            return f"__missing_eid_{row_idx}"
        return t.lower()

    keys = [key_for(out.iloc[i][eid_col], i) for i in range(len(out))]
    out["_eid_pk"] = keys
    n0 = len(out)
    out = out.drop_duplicates(subset=["_eid_pk"], keep="first")
    out = out.drop(columns=["_eid_pk"])
    dropped = n0 - len(out)
    if dropped:
        print(
            f"Deduplicated by {eid_col!r}: dropped {dropped} extra row(s) for the same EID "
            "(kept first occurrence)."
        )
    return out.reset_index(drop=True)


def repeated_participants_table(df):
    """
    Build repeated-participant rows from combined input files before deduplication.
    Returns columns: EID, Name, Repeat Count.
    """
    eid_col = find_eid_column(df)
    if not eid_col:
        return pd.DataFrame(columns=["EID", "Name", "Repeat Count"])

    name_col = find_first_matching_column(df, ["Name", "Student Name", "Full Name"])
    work = df.copy()
    work["_eid_raw"] = work[eid_col].astype(str).str.strip()
    work = work[work["_eid_raw"] != ""]
    work = work[work["_eid_raw"].str.lower() != "nan"]
    if work.empty:
        return pd.DataFrame(columns=["EID", "Name", "Repeat Count"])

    work["_eid_key"] = work["_eid_raw"].str.lower()
    counts = work["_eid_key"].value_counts()
    repeated_keys = counts[counts > 1]
    if repeated_keys.empty:
        return pd.DataFrame(columns=["EID", "Name", "Repeat Count"])

    first_by_key = work.drop_duplicates(subset=["_eid_key"], keep="first").set_index("_eid_key")
    rows = []
    for key, cnt in repeated_keys.items():
        eid_value = first_by_key.at[key, "_eid_raw"] if key in first_by_key.index else key
        if name_col and key in first_by_key.index:
            name_value = str(first_by_key.at[key, name_col]).strip()
            if name_value.lower() == "nan":
                name_value = ""
        else:
            name_value = ""
        rows.append({"EID": eid_value, "Name": name_value, "Repeat Count": int(cnt)})

    out = pd.DataFrame(rows).sort_values(["Repeat Count", "EID"], ascending=[False, True]).reset_index(drop=True)
    return out


_NON_US_CITIZEN_LABEL = "Non - U.S. Citizen"
_NON_US_ABBREV_RE = re.compile(
    r"^non\s*[-–—/]?\s*u\.?\s*s\.?\.?\s*(citizen)?\s*$",
    re.IGNORECASE,
)


def normalize_citizenship(series):
    """
    Normalize missing to Unknown; merge abbreviated Non-U.S. labels into a single
    Non - U.S. Citizen category (e.g. exports split 'Non - U.S.' vs 'Non - U.S. Citizen').
    """
    if series is None:
        return series
    s = normalize_unknown(series)

    def relabel(val):
        if val is None:
            return val
        try:
            if pd.isna(val):
                return val
        except (TypeError, ValueError):
            pass
        t = str(val).strip()
        if t == "Unknown":
            return t
        compact = re.sub(r"[\s\-_.]+", "", t.lower())
        if compact in ("nonus", "nonuscitizen"):
            return _NON_US_CITIZEN_LABEL
        if _NON_US_ABBREV_RE.match(t):
            return _NON_US_CITIZEN_LABEL
        return val

    return s.map(relabel)


def country_distribution_category(df):
    """
    If the participant file has a Country column, return (series, title) for reporting.

    series is normalize_unknown(country) so blanks align with other demographics.
    Column detection matches Gender/Citizenship style via find_first_matching_column.

    Returns None when no Country-like column exists.
    """
    if df is None or df.empty:
        return None
    country_col = find_first_matching_column(
        df,
        [
            "Country",
        ],
    )
    if not country_col:
        return None
    return (normalize_unknown(df[country_col]), "Proportion of Country")


def normalize_unknown(series):
    """Normalize blanks/missing values to 'Unknown' for reporting (one category per field)."""
    s = series.astype(str).str.strip()
    s = s.replace({"": pd.NA})
    unk_tokens = {
        "",
        "nan",
        "none",
        "n/a",
        "na",
        "--",
        "-",
        ".",
        "..",
        "unknown",
        "unk",
        "null",
        "<na>",
    }
    s = s.mask(s.str.lower().isin(unk_tokens), pd.NA)
    return s.fillna("Unknown")

def add_table_to_doc(doc, title, df, style='Table Grid'):
    """Add a formatted Word table with a title."""
    doc.add_paragraph(title, style='Heading 2')
    nrows, ncols = len(df) + 1, len(df.columns)
    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for j, col in enumerate(df.columns):
        cell = table.rows[0].cells[j]
        cell.text = col
        _shade_cell(cell, 'D9E2F3')
    # Data rows
    for i, row in enumerate(df.itertuples(index=False)):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            if df.columns[j] == 'Proportion':
                cell.text = f"{val:.1%}"
            else:
                cell.text = str(val)
    doc.add_paragraph()

def _shade_cell(cell, fill_hex):
    """Apply light shading to a table cell."""
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill_hex)
    cell._tc.get_or_add_tcPr().append(shd)

# Invisible / odd space chars in exports (zero-width, BOM, NBSP) break naive "esl" == checks.
_ESL_STRIP_INVISIBLE = re.compile(r"[\u200b-\u200d\ufeff\u00a0]+")


def _normalized_esl_key(val):
    """Fold Unicode + strip invisibles so ESL, ＥＳＬ, E\u200bSL, etc. match."""
    if val is pd.NA or (isinstance(val, float) and pd.isna(val)):
        return None
    s = unicodedata.normalize("NFKC", str(val))
    s = _ESL_STRIP_INVISIBLE.sub("", s).strip()
    return s.casefold()


# Major / school text that implies ESL for Program Type (word token or full phrase).
_ESL_TOKEN_IN_TEXT = re.compile(r"(?<!\w)esl(?!\w)", re.IGNORECASE)
_ESL_LONG_PHRASE = re.compile(r"english\s+as\s+(?:a\s+)?second\s+language", re.IGNORECASE)


def _text_indicates_esl_program(val):
    """True if free text is ESL or describes English as a Second Language."""
    if val is None:
        return False
    try:
        if pd.isna(val):
            return False
    except (TypeError, ValueError):
        pass
    s = unicodedata.normalize("NFKC", str(val))
    s = _ESL_STRIP_INVISIBLE.sub("", s).strip()
    if not s:
        return False
    if _normalized_esl_key(s) == "esl":
        return True
    if _ESL_TOKEN_IN_TEXT.search(s):
        return True
    cf = re.sub(r"\s+", " ", s.casefold())
    if _ESL_LONG_PHRASE.search(cf):
        return True
    return False


# Advisor Toolkit glitch: euro sign instead of "(E)" before Natural Sciences.
_EURO_PREFIX_NATURAL_SCIENCES = re.compile(
    r"^\s*\u20ac\s*(?=Natural\s+Sciences\b)",
    re.IGNORECASE,
)


def _normalize_euro_natural_sciences_pseudo(val):
    """Map €Natural Sciences → (E)Natural Sciences for pseudo-school parsing."""
    if val is None:
        return val
    try:
        if pd.isna(val):
            return val
    except (TypeError, ValueError):
        pass
    s = unicodedata.normalize("NFKC", str(val))
    return _EURO_PREFIX_NATURAL_SCIENCES.sub("(E)", s)


# Toolkit spelling vs preferred label for Undergraduate Studies bucket display / merges.
_INTERDEPARTMENTAL_LABEL_RE = re.compile(r"\binter[- ]?departmental\b", re.IGNORECASE)


def canonicalize_school_display_name(name):
    """School column: ESL unified; Interdepartmental -> Inter-Department; euro glitch; (X) pseudo-school -> blank."""
    if name is None:
        return name
    try:
        if pd.isna(name):
            return name
    except (TypeError, ValueError):
        pass
    s = str(name).strip()
    if not s:
        return s
    s = str(_normalize_euro_natural_sciences_pseudo(s)).strip()
    if not s:
        return s
    s = _INTERDEPARTMENTAL_LABEL_RE.sub("Inter-Department", s)
    # Toolkit placeholder pseudo-school (X) — treat as no college / blank for counts and merges.
    if re.match(r"^\(\s*X\s*\)", s, re.IGNORECASE):
        return ""
    if _normalized_esl_key(s) == "esl":
        return "ESL"
    m = re.match(r"^\([^)]+\)\s*(.+)$", s)
    if m and _normalized_esl_key(m.group(1).strip()) == "esl":
        return "ESL"
    return s


def _is_option_iii_irregular_label(val):
    """True when Irregular Program is Option III (spacing/punctuation-insensitive); does not match ESL."""
    if val is None:
        return False
    try:
        if pd.isna(val):
            return False
    except (TypeError, ValueError):
        pass
    s = unicodedata.normalize("NFKC", str(val).strip())
    s = s.replace("\xa0", " ").strip().lower()
    if not s or s == "nan":
        return False
    compact = re.sub(r"[\s\-_:]+", "", s)
    return compact in ("optioniii", "option3", "optiii")


def esl_signal_mask(df):
    """
    Per-row ESL signal from Irregular Program cell text, Major, and Pseudo Sch columns.
    Same rules as Program Type (English as a Second Language tokens/phrases).
    """
    irregular_col = next(
        (c for c in df.columns if c.strip().lower() == "irregular program"),
        None,
    )
    irregular_esl_signal = pd.Series(False, index=df.index)
    if irregular_col:
        irregular_esl_signal = df[irregular_col].map(_text_indicates_esl_program)

    major_col = find_first_matching_column(df, ["Maj1 Name", "Major", "Major Name"])
    major_esl = (
        df[major_col].map(_text_indicates_esl_program)
        if major_col
        else pd.Series(False, index=df.index)
    )
    pseudo_cols = [c for c in df.columns if c.strip().lower().startswith("pseudo sch")]
    school_esl = pd.Series(False, index=df.index)
    for c in pseudo_cols:
        school_esl = school_esl | df[c].map(_text_indicates_esl_program)

    return irregular_esl_signal | major_esl | school_esl


def program_type_from_irregular_field(df):
    """
    Program type rules for reporting:
    - Start from Irregular Program: blank -> Regular, else that value (short ESL spellings -> "ESL")
    - Option III is ignored (treated like blank -> Regular/Degree-Seeking unless ESL overrides)
    - Never_Enrolled: ESL (any signal) wins over Unknown; else keep a non-blank Irregular Program value;
      if neither ESL nor other irregular label (still Regular), use Unknown
    - ESL if Irregular Program, Major, or any Pseudo Sch cell indicates ESL / English as a Second Language
      (overrides Regular and other irregular labels when status is not Never_Enrolled)
    - Scholar when Irregular Program indicates scholar / visiting scholar / postdoctoral fellow (not overriding ESL)
    """
    status_col = next(
        (c for c in df.columns if c.strip().lower() == "derived academic status"),
        None,
    )
    irregular_col = next(
        (c for c in df.columns if c.strip().lower() == "irregular program"),
        None,
    )

    # Start as Regular when Irregular Program is blank; Never_Enrolled handled below.
    program = pd.Series(["Regular"] * len(df), index=df.index, dtype=object)

    if irregular_col:
        raw = df[irregular_col].astype(str).str.strip()
        raw = raw.replace({"": pd.NA, "nan": pd.NA, "None": pd.NA, "none": pd.NA})
        raw = raw.mask(raw.map(_is_option_iii_irregular_label), pd.NA)
        esl_keys = raw.map(lambda x: _normalized_esl_key(x) if pd.notna(x) else pd.NA)
        esl_mask = raw.notna() & esl_keys.eq("esl")
        raw = raw.mask(esl_mask, "ESL")
        # If not blank, preserve category value from uploaded file (ESL normalized above).
        program = raw.fillna("Regular")

    is_esl = esl_signal_mask(df)

    never_mask = pd.Series(False, index=df.index)
    if status_col:
        status = df[status_col].astype(str).str.strip().str.lower()
        never_mask = status.eq("never_enrolled")

    # Never_Enrolled: Unknown only if no ESL and no other irregular label (still Regular from blank IP field).
    program.loc[never_mask & ~is_esl & program.eq("Regular")] = "Unknown"
    program.loc[is_esl] = "ESL"
    if irregular_col:
        scholar_ip = df[irregular_col].map(_irregular_program_indicates_scholar)
        program.loc[scholar_ip & ~is_esl] = "Scholar"

    return normalize_unknown(program)


def _compact_for_irregular_scholar_match(val):
    """Normalize Irregular Program text for scholar detection."""
    if val is None:
        return ""
    try:
        if pd.isna(val):
            return ""
    except (TypeError, ValueError):
        pass
    s = str(val).strip().lower()
    if not s or s == "nan":
        return ""
    return re.sub(r"[\s\-_/()]+", "", s)


def _irregular_program_indicates_scholar(val):
    """
    True when Irregular Program maps to Academic Status bucket Scholar:
    scholar, visiting scholar, or postdoctoral fellow (word/phrase match on Irregular Program).
    """
    if val is None:
        return False
    try:
        if pd.isna(val):
            return False
    except (TypeError, ValueError):
        pass
    raw = unicodedata.normalize("NFKC", str(val).strip())
    if not raw or raw.lower() == "nan":
        return False
    compact = _compact_for_irregular_scholar_match(raw)
    if compact == "scholar":
        return True
    if "visitingscholar" in compact:
        return True
    snorm = re.sub(r"\s+", " ", raw.lower()).strip()
    if re.search(r"\bvisiting\s+scholar\b", snorm):
        return True
    if re.search(r"\bscholar\b", snorm):
        return True
    if re.search(
        r"\bpostdoctoral\b(?:\s+fellow)?\b|\bpost[- ]?doctoral\s+fellow\b|\bpost[- ]?doc\s+fellow\b|\bpostdocs?\b",
        snorm,
    ):
        return True
    return False


# AcademicStatus (latest semester column): compact key → report bucket
_LEVEL_GROUP_BY_COMPACT_STATUS = {
    "neverenrolled": "Other",
    "enrlgood": "Undergraduate",
    "gradstud": "Graduate",
    "utbachelorgrad": "Graduate",
    # Historic export typo still seen in some files
    "utbachlorgrad": "Graduate",
}

_ACADEMIC_STATUS_LATEST_SEM_COL_LOWER = (
    "academicstatus (assumed as of student's latest semester)"
)


def _normalize_header_for_match(name):
    """Strip, lower, and unify apostrophes for column-name comparison."""
    n = str(name).strip().lower()
    return n.replace("\u2019", "'").replace("`", "'")


def find_academic_status_latest_semester_column(df):
    """
    Column used for the Graduate / Undergraduate / Other chart:
    AcademicStatus (assumed as of Student's Latest Semester).
    """
    if df is None or len(df.columns) == 0:
        return None
    for c in df.columns:
        if _normalize_header_for_match(c) == _ACADEMIC_STATUS_LATEST_SEM_COL_LOWER:
            return c
    for c in df.columns:
        k = _normalize_header_for_match(c)
        if k.startswith("academicstatus") and "latest" in k and "semester" in k:
            return c
    return None


def _map_academic_status_to_level_group(val):
    """
    Map AcademicStatus (latest semester) to Undergraduate / Graduate / Other.

    Scholar is applied only from Irregular Program (see academic_level_group_series).
    Only Toolkit labels in _LEVEL_GROUP_BY_COMPACT_STATUS are recognized; else Other.
    """
    if val is None:
        return "Other"
    try:
        if pd.isna(val):
            return "Other"
    except (TypeError, ValueError):
        pass
    s = str(val).strip().lower().lstrip("\ufeff")
    s = unicodedata.normalize("NFKC", s)
    s = s.replace("\xa0", " ").strip()
    for ch in ("\u2013", "\u2014", "\u2212"):
        s = s.replace(ch, "-")
    if s in {"", "nan", "none", "n/a", "na", "--", "-"}:
        return "Other"
    compact = re.sub(r"[\s\-_/()]+", "", s)
    return _LEVEL_GROUP_BY_COMPACT_STATUS.get(compact, "Other")


def _infer_level_from_career_like(val):
    """
    Map free-text career fields to Undergraduate / Graduate when obvious.
    Returns pd.NA when unknown.
    """
    if val is None:
        return pd.NA
    try:
        if pd.isna(val):
            return pd.NA
    except (TypeError, ValueError):
        pass
    s = unicodedata.normalize("NFKC", str(val).strip())
    s = s.replace("\xa0", " ").strip().lower()
    if s in {"", "nan", "none", "n/a", "na", "--", "-"}:
        return pd.NA
    if re.search(r"\bundergraduate\b|\bundergrad\b|\bugrd\b", s):
        return "Undergraduate"
    if re.search(
        r"\bgraduate\b|\bgrad\s+stud|\bmasters?\b|\bdoctoral\b|\bph\.?\s*d\b|\bjd\b|\bmd\b",
        s,
    ):
        return "Graduate"
    return pd.NA


def _map_derived_status_to_undergrad_grad(val):
    """
    Map Derived Academic Status to Undergraduate / Graduate when unambiguous.
    Returns pd.NA when not an enrollment-level signal (e.g. Never_Enrolled).
    """
    if val is None:
        return pd.NA
    try:
        if pd.isna(val):
            return pd.NA
    except (TypeError, ValueError):
        pass
    s = str(val).strip().lower().lstrip("\ufeff")
    s = unicodedata.normalize("NFKC", s)
    s = s.replace("\xa0", " ").strip()
    for ch in ("\u2013", "\u2014", "\u2212"):
        s = s.replace(ch, "-")
    if s in {"", "nan", "none", "n/a", "na", "--", "-"}:
        return pd.NA
    compact = re.sub(r"[\s\-_/]+", "", s)
    if compact == "enrlgood":
        return "Undergraduate"
    if compact in ("gradstud", "utbachelorgrad", "utbachlorgrad"):
        return "Graduate"
    return pd.NA


def academic_level_group_series(df):
    """
    Return a Series of 'Undergraduate' | 'Graduate' | 'Scholar' | 'ESL' | 'Other' aligned to df.index,
    or None if AcademicStatus (latest semester) is not present.

    Scholar is assigned only when Irregular Program indicates scholar, visiting scholar,
    or postdoctoral fellow (see _irregular_program_indicates_scholar), and ESL does not apply.

    ESL uses the same Irregular Program / Major / Pseudo Sch signals as Program Type and overrides
    other buckets so ESL is separated from Other.

    When Irregular Program is Option III and AcademicStatus maps to Other, level is inferred
    from Career-like columns first, then Derived Academic Status (Advisor Toolkit fields
    that appear before Irregular Program in typical exports).
    """
    status_col = find_academic_status_latest_semester_column(df)
    if not status_col:
        return None
    is_esl = esl_signal_mask(df)
    out = df[status_col].map(_map_academic_status_to_level_group)

    irregular_col = next(
        (c for c in df.columns if c.strip().lower() == "irregular program"),
        None,
    )
    if irregular_col:
        scholar_ip = df[irregular_col].map(_irregular_program_indicates_scholar)
        scholar_rows = scholar_ip & ~is_esl
        if scholar_rows.any():
            out = out.mask(scholar_rows, "Scholar")

        opt3 = df[irregular_col].map(_is_option_iii_irregular_label)
        fc = opt3 & out.eq("Other")
        if fc.any():
            career_col = find_first_matching_column(
                df,
                ["Career", "Academic Career", "Student Career", "Career Description"],
            )
            das_col = next(
                (c for c in df.columns if c.strip().lower() == "derived academic status"),
                None,
            )
            sub = out.loc[fc].copy()
            if career_col:
                inf_c = df.loc[fc, career_col].map(_infer_level_from_career_like)
                sub = inf_c.where(inf_c.notna(), sub)
            if das_col:
                inf_d = df.loc[fc, das_col].map(_map_derived_status_to_undergrad_grad)
                sub = inf_d.where(inf_d.notna(), sub)
            out.loc[fc] = sub

    if is_esl.any():
        out.loc[is_esl] = "ESL"

    return out


def ordered_level_counts(level_series, level_order=("Undergraduate", "Graduate", "Scholar", "ESL", "Other")):
    """value_counts restricted to level_order; drops categories with zero count."""
    counts = level_series.value_counts()
    ordered = pd.Series({k: int(counts.get(k, 0)) for k in level_order})
    return ordered[ordered > 0]


# Embedded school lookup: script looks for this file in script dir, CSV dir, or Downloads
SCHOOL_LOOKUP_FILENAME = "COLA Toolkit, Spring 2026.csv"
# Enrollment reference for representation comparison (optional)
ENROLLMENT_REFERENCE_FILENAME = "All_International_Students_Enrolled.csv"
COUNTRY_ENROLLMENT_REFERENCE_FILENAME = "All_International_Students_By_Country.csv"

# Built‑in code → school mapping (used if no external lookup file is present).
DEFAULT_SCHOOL_CODE_LOOKUP = {
    "2": "Business Administration",
    "3": "Education",
    "4": "Engineering",
    "5": "Fine Arts",
    "7": "Law School",
    "8": "Pharmacy",
    "9": "Architecture",
    "B": "Graduate Business",
    "C": "Communication",
    "E": "Natural Sciences",
    "F": "Civic Leadership",
    "J": "Geosciences",
    "L": "Liberal Arts",
    "M": "Medical School",
    "N": "Nursing",
    "P": "Information",
    "S": "Social Work",
    "T": "Public Affairs",
    "U": "Undergraduate Studies",
}


def resolve_school_lookup_path(csv_dir, script_dir=None):
    """Return path to school lookup file if it exists in script dir, CSV dir, or user Downloads."""
    candidates = []
    if script_dir:
        candidates.append(os.path.join(script_dir, SCHOOL_LOOKUP_FILENAME))
    candidates.append(os.path.join(csv_dir, SCHOOL_LOOKUP_FILENAME))
    downloads = os.path.join(os.path.expanduser("~"), "Downloads", SCHOOL_LOOKUP_FILENAME)
    candidates.append(downloads)
    for p in candidates:
        if os.path.isfile(p):
            return p
    return None


def resolve_enrollment_path(csv_dir, script_dir=None, explicit_path=None):
    """Return path to enrollment reference CSV if provided or found (csv_dir, script_dir, Downloads)."""
    if explicit_path and os.path.isfile(explicit_path):
        return explicit_path
    candidates = []
    if script_dir:
        candidates.append(os.path.join(script_dir, ENROLLMENT_REFERENCE_FILENAME))
    candidates.append(os.path.join(csv_dir, ENROLLMENT_REFERENCE_FILENAME))
    downloads = os.path.join(os.path.expanduser("~"), "Downloads", ENROLLMENT_REFERENCE_FILENAME)
    candidates.append(downloads)
    for p in candidates:
        if os.path.isfile(p):
            return p
    return None


def resolve_country_enrollment_path(csv_dir, script_dir=None):
    """Return path to international enrollment-by-country reference CSV if found (script dir, csv dir, Downloads)."""
    candidates = []
    if script_dir:
        candidates.append(os.path.join(script_dir, COUNTRY_ENROLLMENT_REFERENCE_FILENAME))
    candidates.append(os.path.join(csv_dir, COUNTRY_ENROLLMENT_REFERENCE_FILENAME))
    downloads = os.path.join(os.path.expanduser("~"), "Downloads", COUNTRY_ENROLLMENT_REFERENCE_FILENAME)
    candidates.append(downloads)
    for p in candidates:
        if os.path.isfile(p):
            return p
    return None


def _normalize_country_key(name):
    """Normalize country labels so enrollment and participant exports match when spelling differs slightly."""
    if name is None:
        return ""
    try:
        if pd.isna(name):
            return ""
    except (TypeError, ValueError):
        pass
    s = unicodedata.normalize("NFKC", str(name)).strip()
    for ch in ("\u2018", "\u2019", "\u201b", "\u2032"):
        s = s.replace(ch, "'")
    s = re.sub(r"\s+", " ", s)
    return s.casefold()


def load_enrollment_by_country(path):
    """
    Load international enrollment counts by country from a summary CSV.
    Expects a 'Country' column and one of Count, Enrollment, or Students.
    Optional 'Continent' column is returned for display in comparison tables.
    Ignores blank country rows and summary rows labeled Grand Total.

    Returns
    -------
    tuple
        (counts Series keyed by normalized country, total enrollment, continent_by_key, display_name_by_key)
    """
    try:
        enr = pd.read_csv(path)
    except Exception:
        return pd.Series(dtype=float), 0, {}, {}
    enr.columns = enr.columns.str.strip()
    if "Country" not in enr.columns:
        return pd.Series(dtype=float), 0, {}, {}
    count_col = next(
        (c for c in ("Count", "Enrollment", "Students") if c in enr.columns),
        None,
    )
    if not count_col:
        return pd.Series(dtype=float), 0, {}, {}
    has_continent = "Continent" in enr.columns
    enr = enr.dropna(subset=["Country"])
    enr["Country"] = enr["Country"].astype(str).str.strip()
    enr = enr[enr["Country"] != ""]
    enr = enr[~enr["Country"].str.casefold().eq("grand total")]
    counts_by_key = {}
    label_by_key = {}
    continent_by_key = {}
    for _, row in enr.iterrows():
        raw = str(row["Country"]).strip()
        key = _normalize_country_key(raw)
        if not key:
            continue
        try:
            c = int(float(row[count_col]))
        except (TypeError, ValueError):
            continue
        counts_by_key[key] = counts_by_key.get(key, 0) + c
        if key not in label_by_key:
            label_by_key[key] = raw
        if has_continent and key not in continent_by_key:
            ct = row["Continent"]
            if ct is not None and str(ct).strip() and str(ct).strip().lower() != "nan":
                continent_by_key[key] = str(ct).strip()
    if not counts_by_key:
        return pd.Series(dtype=float), 0, {}, {}
    total = int(sum(counts_by_key.values()))
    counts = pd.Series(counts_by_key, dtype=float)
    return counts, total, continent_by_key, label_by_key


def build_country_representation_comparison(
    participation_counts,
    total_participants,
    enrollment_counts,
    total_enrollment,
    enrollment_display_by_key=None,
    continent_by_key=None,
):
    """
    Same logic as school representation: Enrollment %, Participation %, ratio = part% / enr%.
    participation_counts: value_counts–style Series (index = country label as in participant file).
    enrollment_counts: Series indexed by _normalize_country_key (as from load_enrollment_by_country).
    enrollment_display_by_key: optional dict norm_key -> preferred display name (e.g. from reference file).
    continent_by_key: optional dict norm_key -> continent label from reference file.
    """
    empty_cols = [
        "Continent",
        "Country",
        "Enrollment Count",
        "Enrollment %",
        "Participant EIDs",
        "Participation %",
        "Representation Ratio",
    ]
    if total_participants <= 0 or total_enrollment <= 0:
        return pd.DataFrame(columns=empty_cols)

    part_norm = {}
    part_label = {}
    for country, val in participation_counts.items():
        key = _normalize_country_key(country)
        if not key:
            continue
        part_norm[key] = part_norm.get(key, 0) + int(val)
        if key not in part_label:
            part_label[key] = str(country).strip()

    enr_norm = {}
    for k, val in enrollment_counts.items():
        key = _normalize_country_key(k)
        if not key:
            continue
        enr_norm[key] = enr_norm.get(key, 0) + int(val)

    all_keys = sorted(set(part_norm.keys()) | set(enr_norm.keys()))
    rows = []
    disp = enrollment_display_by_key or {}
    cont = continent_by_key or {}
    for key in all_keys:
        enc = int(enr_norm.get(key, 0))
        prc = int(part_norm.get(key, 0))
        if enc == 0 and prc == 0:
            continue
        enr_pct = (enc / total_enrollment) * 100 if total_enrollment > 0 else 0.0
        part_pct = (prc / total_participants) * 100 if total_participants > 0 else 0.0
        ratio = (part_pct / enr_pct) if enr_pct else float("nan")
        display_country = disp.get(key) or part_label.get(key) or key
        continent_val = cont.get(key, "")
        rows.append(
            {
                "Continent": continent_val,
                "Country": display_country,
                "Enrollment Count": enc,
                "Enrollment %": enr_pct,
                "Participant EIDs": prc,
                "Participation %": part_pct,
                "Representation Ratio": ratio,
            }
        )
    out = pd.DataFrame(rows)
    if not out.empty:
        out = out.assign(_sortkey=out["Country"].astype(str).str.casefold()).sort_values("_sortkey").drop(
            columns=["_sortkey"]
        ).reset_index(drop=True)
        if out["Continent"].astype(str).str.strip().eq("").all():
            out = out.drop(columns=["Continent"])
    return out


def load_school_lookup(path):
    """Load school code -> official school name mapping from a lookup CSV, layered over defaults."""
    mapping = dict(DEFAULT_SCHOOL_CODE_LOOKUP)
    if not path:
        return mapping
    try:
        lut = pd.read_csv(path)
    except FileNotFoundError:
        return mapping

    lut.columns = lut.columns.str.strip()
    if not {"Code", "School"}.issubset(lut.columns):
        print(f"School lookup file {path} is missing 'Code' or 'School' columns. Using built‑in school code mapping.")
        return mapping

    lut = lut[
        (lut["Code"].astype(str).str.strip() != "")
        & (lut["School"].astype(str).str.strip() != "")
    ]
    csv_mapping = dict(
        zip(
            lut["Code"].astype(str).str.strip(),
            lut["School"].astype(str).str.strip(),
        )
    )
    mapping.update(csv_mapping)
    return mapping


def translate_pseudo_school(series, code_to_school):
    """Translate '(Code)School' pseudo values to official school names using lookup."""
    if series is None or not code_to_school:
        return series

    def _translate(val):
        if pd.isna(val):
            return val
        s = _normalize_euro_natural_sciences_pseudo(str(val))
        # Extract code inside parentheses, e.g. "(E)Natural Sciences" -> "E"
        m = re.search(r"\(([^)]+)\)", s)
        if not m:
            return s
        code = m.group(1).strip()
        if _pseudo_code_is_placeholder_x(code):
            return ""
        return code_to_school.get(code, s)

    return series.apply(_translate)


def _pseudo_code_is_graduate_school(code_str):
    """Toolkit pseudo-school code 6 is Graduate School (dropped from participation/enrollment school lists)."""
    c = str(code_str).strip()
    return c.isdigit() and int(c) == 6


def _pseudo_code_is_placeholder_x(code_str):
    """Toolkit pseudo-school code X means blank / no school for reporting."""
    return str(code_str).strip().upper() == "X"


def _is_graduate_school(school_name):
    """True if this bucket is UT pseudo-school code 6 / Graduate School (excluded from school breakdown)."""
    if school_name is None:
        return False
    try:
        if pd.isna(school_name):
            return False
    except (TypeError, ValueError):
        pass
    raw = str(school_name).strip()
    if not raw:
        return False
    low = raw.lower()
    if low == "graduate school":
        return True
    m = re.match(r"^\(([^)]+)\)\s*(.*)$", raw)
    if m:
        if _pseudo_code_is_graduate_school(m.group(1)):
            return True
        rest = m.group(2).strip().lower()
        if rest == "graduate school":
            return True
    return False


def _parse_schools_from_cell(val, code_to_school):
    """Parse one cell (may contain 'A/ B'). Return set of translated school names."""
    if pd.isna(val) or str(val).strip() == "":
        return set()
    parts = [p.strip() for p in str(val).split("/") if p.strip()]
    result = set()
    for s in parts:
        s = _normalize_euro_natural_sciences_pseudo(s)
        m = re.search(r"\(([^)]+)\)", s)
        if m:
            code = m.group(1).strip()
            if _pseudo_code_is_graduate_school(code):
                continue
            if _pseudo_code_is_placeholder_x(code):
                continue
            name = canonicalize_school_display_name(code_to_school.get(code, s))
            if name and str(name).strip() and not _is_graduate_school(name):
                result.add(name)
        else:
            name = canonicalize_school_display_name(s)
            if name and str(name).strip() and not _is_graduate_school(name):
                result.add(name)
    return result


def student_based_school_proportions(df, code_to_school):
    """
    College/School proportions: denominator is len(df) = one row per unique EID.
    An EID in multiple schools (e.g. Graduate + COLA) is counted in each school;
    proportions are (EIDs in school X) / total unique EIDs, so they can sum to more than 100%.
    Uses Pseudo Sch1 and Pseudo Sch2; cells with 'A/ B' are split into multiple schools.
    Returns (counts_series, proportions_df) for table/chart.
    """
    total_students = len(df)
    if total_students == 0:
        return pd.Series(dtype=object), pd.DataFrame(columns=["Category", "Count", "Proportion"])

    # Be flexible about pseudo school column names (e.g. 'Pseudo Sch1', 'Pseudo Sch Maj1', etc.).
    pseudo_cols = [
        c for c in df.columns if c.strip().lower().startswith("pseudo sch")
    ]
    if not pseudo_cols:
        return pd.Series(dtype=object), pd.DataFrame(columns=["Category", "Count", "Proportion"])

    col1 = pseudo_cols[0]
    col2 = pseudo_cols[1] if len(pseudo_cols) > 1 else None
    student_school_counts = Counter()
    for idx, row in df.iterrows():
        schools = _parse_schools_from_cell(row[col1], code_to_school)
        if col2:
            schools |= _parse_schools_from_cell(row[col2], code_to_school)
        valid_schools = [s for s in schools if s and not _is_graduate_school(s)]
        for s in valid_schools:
            if s and not _is_graduate_school(s):
                student_school_counts[s] += 1
        if not valid_schools:
            student_school_counts["Unknown"] += 1

    if not student_school_counts:
        return pd.Series(dtype=object), pd.DataFrame(columns=["Category", "Count", "Proportion"])

    counts = pd.Series(student_school_counts)
    proportions = counts / total_students
    proportions_df = pd.DataFrame({
        "Category": counts.index.astype(str),
        "Count": counts.values,
        "Proportion": proportions.values,
    })
    return counts, proportions_df


def load_enrollment_by_school(path, code_to_school):
    """
    Load enrollment counts by school from a CSV.
    Supports: (1) Student-level CSV with Pseudo Sch1 / Pseudo Sch2 (same as event data);
              (2) Summary CSV with 'School' and 'Enrollment' or 'Count'.
    Returns (series: school -> count, total_enrollment).
    """
    try:
        enr = pd.read_csv(path)
    except Exception:
        return pd.Series(dtype=object), 0
    enr.columns = enr.columns.str.strip()
    # Summary format: School + Enrollment, Count, or Students
    if "School" in enr.columns:
        count_col = next(
            (c for c in ("Enrollment", "Count", "Students") if c in enr.columns),
            None,
        )
        if count_col:
            enr = enr.dropna(subset=["School"])
            enr["School"] = enr["School"].map(canonicalize_school_display_name)
            enr = enr[enr["School"].astype(str).str.strip() != ""]
            enr = enr[~enr["School"].map(_is_graduate_school)]
            total_enrollment = int(enr[count_col].sum())
            counts = enr.groupby("School", as_index=True)[count_col].sum()
            return counts, total_enrollment
    # Student-level format (same as event CSV)
    if "Pseudo Sch1" in enr.columns:
        counts_series, _ = student_based_school_proportions(enr, code_to_school)
        total = len(enr)
        return counts_series, total
    return pd.Series(dtype=object), 0


def build_representation_comparison(participation_counts, total_participants, enrollment_counts, total_enrollment):
    """
    Build comparison table: Enrollment %, Participation %, Representation Ratio.
    Ratio = (Participation %) / (Enrollment %). ≈1 proportional, >1 overrepresented, <1 underrepresented.
    total_participants must be the count of unique EIDs in the event file (same denominator as school proportions).

    If a school appears twice with and without a leading code, e.g.
    "(2)Business Administration" and "Business Administration", they are
    merged into a single row using the plain-school label.
    """
    if total_participants <= 0 or total_enrollment <= 0:
        return pd.DataFrame(
            columns=[
                "School",
                "Enrollment Count",
                "Enrollment %",
                "Participant EIDs",
                "Participation %",
                "Representation Ratio",
            ]
        )

    def _normalize_label(label: str) -> str:
        raw = str(label).strip()
        s = str(_normalize_euro_natural_sciences_pseudo(raw)).strip()
        m = re.match(r"\(([^)]+)\)\s*(.+)", s)
        core = m.group(2).strip() if m else s
        return canonicalize_school_display_name(core)

    # Normalize and merge participation counts
    part_norm = {}
    for school, val in participation_counts.items():
        key = _normalize_label(school)
        part_norm[key] = part_norm.get(key, 0) + int(val)

    # Normalize and merge enrollment counts
    enr_norm = {}
    for school, val in enrollment_counts.items():
        key = _normalize_label(school)
        enr_norm[key] = enr_norm.get(key, 0) + int(val)

    all_schools = sorted(set(part_norm.keys()) | set(enr_norm.keys()))
    rows = []
    for school in all_schools:
        enc = int(enr_norm.get(school, 0))
        prc = int(part_norm.get(school, 0))
        # Skip rows where both enrollment and participation are zero
        if enc == 0 and prc == 0:
            continue
        enr_pct = (enc / total_enrollment) * 100 if total_enrollment > 0 else 0.0
        part_pct = (prc / total_participants) * 100 if total_participants > 0 else 0.0
        ratio = (part_pct / enr_pct) if enr_pct else float("nan")
        rows.append(
            {
                "School": school,
                "Enrollment Count": enc,
                "Enrollment %": enr_pct,
                "Participant EIDs": prc,
                "Participation %": part_pct,
                "Representation Ratio": ratio,
            }
        )
    return pd.DataFrame(rows)


def comparison_table_to_doc(doc, df, intro_paragraph=None):
    """Add representation comparison table to document."""
    if df.empty:
        return
    if intro_paragraph is None:
        intro_paragraph = (
            "Enrollment % = (School enrollment / total enrollment in reference file) × 100. "
            "Participation % = (participant EIDs in school / total unique event EIDs) × 100. "
            "Ratio = Participation % / Enrollment % (≈1 proportional, >1 overrepresented, <1 underrepresented)."
        )
    doc.add_paragraph(intro_paragraph, style="Normal")
    nrows, ncols = len(df) + 1, len(df.columns)
    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, col in enumerate(df.columns):
        cell = table.rows[0].cells[j]
        cell.text = col
        _shade_cell(cell, "D9E2F3")
    for i, row in enumerate(df.itertuples(index=False)):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            if "Ratio" in df.columns[j]:
                cell.text = f"{val:.2f}" if pd.notna(val) else "—"
            elif "%" in df.columns[j]:
                cell.text = f"{val:.1f}%"
            else:
                cell.text = str(val)
    doc.add_paragraph()


def side_by_side_bar_chart_bytes(
    comparison_df,
    title="Enrollment % vs Participation % by School",
    label_column="School",
):
    """Side-by-side bar chart: Enrollment % and Participation % per category (school or country)."""
    if comparison_df.empty or label_column not in comparison_df.columns:
        buf = io.BytesIO()
        plt.figure(figsize=(8, 4))
        plt.text(0.5, 0.5, "No data", ha="center", va="center")
        plt.savefig(buf, format="png", dpi=120, bbox_inches="tight")
        plt.close()
        buf.seek(0)
        return buf
    df = comparison_df.sort_values("Enrollment %", ascending=False).head(16)
    x = range(len(df))
    w = 0.35
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.bar([i - w / 2 for i in x], df["Enrollment %"], width=w, label="Enrollment %", color="steelblue")
    ax.bar([i + w / 2 for i in x], df["Participation %"], width=w, label="Participation %", color="coral")
    ax.set_xticks(x)
    ax.set_xticklabels(
        df[label_column].astype(str).str[:30]
        + df[label_column].astype(str).str.len().gt(30).map({True: "…", False: ""}),
        rotation=45,
        ha="right",
    )
    ax.set_ylabel("Percentage")
    ax.set_title(title)
    ax.legend()
    ax.axhline(y=0, color="gray", linewidth=0.5)
    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format="png", dpi=120, bbox_inches="tight")
    plt.close()
    buf.seek(0)
    return buf


def representation_ratio_chart_bytes(
    comparison_df,
    title="Representation Ratio & Over/Under by School",
    label_column="School",
):
    """Bar chart of representation ratio (1 = proportional). Gray ≈ proportional, red under, blue/green over."""
    if comparison_df.empty or "Representation Ratio" not in comparison_df.columns:
        buf = io.BytesIO()
        plt.figure(figsize=(8, 4))
        plt.text(0.5, 0.5, "No data", ha="center", va="center")
        plt.savefig(buf, format="png", dpi=120, bbox_inches="tight")
        plt.close()
        buf.seek(0)
        return buf
    df = comparison_df.dropna(subset=["Representation Ratio"]).sort_values("Representation Ratio", ascending=True).tail(16)
    if df.empty:
        buf = io.BytesIO()
        plt.figure(figsize=(8, 4))
        plt.savefig(buf, format="png", dpi=120, bbox_inches="tight")
        plt.close()
        buf.seek(0)
        return buf
    colors = ["gray" if 0.98 <= r <= 1.02 else "tomato" if r < 1 else "forestgreen" for r in df["Representation Ratio"]]
    fig, ax = plt.subplots(figsize=(10, 5))
    y_pos = range(len(df))
    ax.barh(y_pos, df["Representation Ratio"], color=colors)
    ax.axvline(x=1.0, color="black", linestyle="--", linewidth=1, label="Proportional (1.0)")
    ax.set_yticks(y_pos)
    ax.set_yticklabels(
        df[label_column].astype(str).str[:35]
        + df[label_column].astype(str).str.len().gt(35).map({True: "…", False: ""})
    )
    ax.set_xlabel("Representation Ratio (gray ≈ proportional, green over, red under)")
    ax.set_title(title)
    ax.legend()
    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format="png", dpi=120, bbox_inches="tight")
    plt.close()
    buf.seek(0)
    return buf


def _prepare_pie_counts(counts, max_slices=12):
    """Group small slices into 'Other' if there are too many."""
    if len(counts) <= max_slices:
        return counts
    top = counts.head(max_slices - 1)
    other_count = counts.iloc[max_slices - 1:].sum()
    other_label = f"Other ({len(counts) - max_slices + 1} categories)"
    return pd.concat([top, pd.Series({other_label: other_count})])

def pie_chart_to_bytes(counts, title, max_slices=12):
    """Draw a pie chart with legend (no labels on slices) to avoid overlapping text."""
    counts = _prepare_pie_counts(counts, max_slices)
    labels = [str(x)[:40] + ('...' if len(str(x)) > 40 else '') for x in counts.index]
    fig, ax = plt.subplots(figsize=(7, 5))
    wedges, texts, autotexts = ax.pie(
        counts.values,
        labels=None,
        autopct='%1.1f%%',
        startangle=90,
        pctdistance=0.6,
        explode=[0.02] * len(counts),
    )
    ax.set_title(title, fontsize=12)
    ax.legend(wedges, labels, title='Category', loc='center left', bbox_to_anchor=(1, 0.5), fontsize=8)
    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=120, bbox_inches='tight')
    plt.close()
    buf.seek(0)
    return buf

def clean_unknown_students(df):
    """Remove invalid student rows (empty/missing names). Keep Never_Enrolled rows (Program Type uses ESL / Irregular / Unknown rules)."""
    n_before = len(df)
    # Drop rows with no name (empty or blank)
    if "Name" in df.columns:
        df = df[df["Name"].astype(str).str.strip().str.lower() != "nan"]
        df = df[df["Name"].astype(str).str.strip() != ""]
    n_after = len(df)
    removed = n_before - n_after
    if removed > 0:
        print(f"Cleaning: removed {removed} row(s) with missing student identity. Analyzed {n_after} students.")
    return df.reset_index(drop=True)

def _parse_never_enrolled_eids(raw_text):
    """
    Parse Advisor Toolkit messages like
    'dk33895 does not appear to have ever enrolled.'
    and return a sorted list of unique EIDs.
    """
    if not raw_text:
        return []
    eids = set()
    for line in str(raw_text).splitlines():
        line = line.strip()
        if not line:
            continue
        # Grab the first token before whitespace or '('
        m = re.match(r"([A-Za-z0-9]+)", line)
        if m:
            eids.add(m.group(1))
    return sorted(eids)


def _flatten_event_csv_paths(obj):
    """
    Yield filesystem path strings from a str, PathLike, or nested iterable of those.

    str is not treated as iterable (so we do not split into characters). This avoids
    passing a nested list into pd.read_csv(), which raises:
    ValueError: Invalid file path or buffer object type: <class 'list'>.
    """
    if isinstance(obj, str):
        yield obj
        return
    if isinstance(obj, os.PathLike):
        yield os.fspath(obj)
        return
    if isinstance(obj, (bytes, bytearray)):
        yield os.fspath(obj)
        return
    try:
        it = iter(obj)
    except TypeError as e:
        raise TypeError(
            f"Expected a path or iterable of paths, got {type(obj).__name__}"
        ) from e
    for item in it:
        yield from _flatten_event_csv_paths(item)


def _normalize_event_csv_paths(event_csv_path):
    """Return a list of path strings (str or PathLike -> single-element list; nested lists flattened)."""
    if event_csv_path is None:
        return []
    return [p for p in _flatten_event_csv_paths(event_csv_path) if p]


def generate_report(event_csv_path, enrollment_reference_path=None, never_enrolled_notes=None):
    """
    Generate the Word report for one or more event participants CSVs (combined into one report).

    Parameters
    ----------
    event_csv_path : str, os.PathLike, or sequence of those
        Path to the event participants CSV, or several paths. Rows from all
        files are concatenated in order, then cleaned and deduplicated by EID
        (first occurrence wins across files).
    enrollment_reference_path : str or None, optional
        Optional path to the enrollment reference CSV. If None, the script
        will look for All_International_Students_Enrolled.csv in the event
        CSV directory, script directory, or user Downloads.
    never_enrolled_notes : str or None, optional
        Optional raw text from Advisor Toolkit listing EIDs that
        "do not appear to have ever enrolled". Those EIDs are not rows in
        the participant CSV; the report describes them separately and counts
        each only in the irregular-program environment breakdown, not in N
        for the other tables/charts (which use unique EIDs from the file).

    Returns
    -------
    str
        Path to the generated .docx report.
    """
    script_dir = os.path.dirname(os.path.abspath(sys.argv[0])) if sys.argv else os.getcwd()
    event_paths = _normalize_event_csv_paths(event_csv_path)
    if not event_paths:
        raise ValueError("At least one event participants CSV path is required.")

    frames = []
    for p in event_paths:
        try:
            chunk = pd.read_csv(p)
        except FileNotFoundError:
            print(f"File not found: {p}")
            raise
        chunk.columns = chunk.columns.str.strip()
        frames.append(chunk)
    df = pd.concat(frames, ignore_index=True)

    # Clean before any analysis
    df = clean_unknown_students(df)
    repeated_df = repeated_participants_table(df)
    df = dedupe_participants_by_eid(df)
    if len(df) == 0:
        raise ValueError("No rows remaining after cleaning; cannot generate report.")
    n_unique_eids = len(df)

    # Parse "never enrolled" EIDs from optional notes
    never_enrolled_eids = _parse_never_enrolled_eids(never_enrolled_notes)

    # School proportions: denominator = unique EIDs (one row per EID; multi-school cells count in each school)
    primary_path = event_paths[0]
    csv_dir = os.path.dirname(os.path.abspath(primary_path))
    school_lookup_path = resolve_school_lookup_path(csv_dir, script_dir)
    code_to_school = load_school_lookup(school_lookup_path) if school_lookup_path else {}
    school_counts, school_proportions_df = student_based_school_proportions(df, code_to_school)

    country_col = find_first_matching_column(df, ["Country"])
    country_counts = None
    if country_col:
        country_counts = normalize_unknown(df[country_col]).value_counts()

    if len(event_paths) == 1:
        base = os.path.splitext(os.path.basename(event_paths[0]))[0]
    else:
        stem0 = os.path.splitext(os.path.basename(primary_path))[0]
        base = f"{stem0}_merged_{len(event_paths)}files"
    out_dir = os.path.dirname(os.path.abspath(primary_path))
    out_docx = os.path.join(out_dir, f"{base}_report.docx")

    doc = Document()
    doc.add_heading('Event Participant Proportions Report', 0)
    if len(event_paths) == 1:
        doc.add_paragraph(f"Source: {event_paths[0]}")
    else:
        doc.add_paragraph("Sources (combined in this order):")
        for p in event_paths:
            doc.add_paragraph(f"• {p}")
    file_phrase = "file" if len(event_paths) == 1 else "files (concatenated in order, then deduplicated)"
    doc.add_paragraph(
        f"This report is based on N = {n_unique_eids} unique EID(s) from the event participant {file_phrase} (not on enrollment "
        "status). Duplicate rows for the same EID are dropped; the kept row is the first across the combined data. "
        "Rows with missing names are excluded before deduplication. Missing demographic values are shown as Unknown. "
        "For Program Type (shown as Degree-Seeking vs ESL, Scholar, Unknown, Irregular in tables/charts), rows with Derived Academic Status "
        "Never_Enrolled are Unknown only when they have no ESL signal (major/school/Irregular Program) and no other Irregular Program "
        "label; ESL is prioritized over Scholar and Unknown. Scholar matches the same Irregular Program rules as Academic Status. "
        "Irregular Program Option III is ignored for Program Type (Degree-Seeking unless ESL). "
        "For Academic Status, Scholar is assigned only when Irregular Program indicates scholar, visiting scholar, or postdoctoral fellow "
        "(and ESL does not apply); ESL is its own bucket using the same Irregular Program / major / pseudo-school signals as Program Type. "
        "Option III rows that would otherwise be Other use Career or Derived Academic Status for Undergraduate vs Graduate."
    )

    # Data checks from the full event-series input (all uploaded files combined).
    doc.add_heading("Data checks", level=1)
    if repeated_df.empty:
        doc.add_paragraph(
            "No repeated participants were found across the combined event-series input files."
        )
    else:
        doc.add_paragraph(
            "Repeated participants across the combined event-series input "
            "(counts are before EID deduplication)."
        )
        add_table_to_doc(
            doc,
            f"Repeated participants (N repeated EIDs = {len(repeated_df)})",
            repeated_df,
        )

    # Build list of categories from the uploaded participant CSV only.
    categories = []

    # Explicit demographic charts requested: Gender and Citizenship
    gender_col = find_first_matching_column(df, ["Gender"])
    if gender_col:
        categories.append((normalize_unknown(df[gender_col]), "Proportion of Gender"))

    citizenship_col = find_first_matching_column(df, ["Citizenship"])
    if citizenship_col:
        categories.append((normalize_unknown(normalize_citizenship(df[citizenship_col])), "Proportion of Citizenship"))

    country_cat = country_distribution_category(df)
    if country_cat:
        categories.append(country_cat)

    # Irregular = something in 'Irregular Program' field (student from IP)
    df["Program Type"] = program_type_from_irregular_field(df)
    # For program type proportions, also count each "never enrolled" EID as Irregular in the
    # environment. They are only added to this Regular vs Irregular breakdown, not to the
    # other demographic tables/charts.
    if never_enrolled_eids:
        program_series = pd.concat(
            [
                df["Program Type"],
                pd.Series(
                    ["Irregular"] * len(never_enrolled_eids),
                    name="Program Type",
                ),
            ],
            ignore_index=True,
        )
    else:
        program_series = df["Program Type"]
    program_for_report = normalize_unknown(program_series).replace({"Regular": "Degree-Seeking"})
    categories.append(
        (
            program_for_report,
            "Proportion of Program Type",
        )
    )

    level_series = academic_level_group_series(df)

    # Add formatted tables (one per category), with explicit sample sizes
    doc.add_heading('Summary tables', level=1)
    if not school_proportions_df.empty:
        add_table_to_doc(
            doc,
            f"Proportion of College/School (unique EIDs, N = {n_unique_eids})",
            school_proportions_df,
        )
    for series, title in categories:
        # Sample size for this category: non-missing, non-blank entries
        valid = series.dropna().astype(str).str.strip()
        valid_n = (valid != "").sum()
        titled = f"{title} (N = {valid_n})"
        counts, proportions, tbl = get_proportions_df(series)
        add_table_to_doc(doc, titled, tbl)

    if level_series is not None:
        level_counts = ordered_level_counts(level_series)
        if not level_counts.empty:
            total_lv = int(level_counts.sum())
            tbl_lv = pd.DataFrame(
                {
                    "Category": level_counts.index.astype(str),
                    "Count": level_counts.values,
                    "Proportion": (level_counts.values / total_lv).astype(float),
                }
            )
            add_table_to_doc(
                doc,
                f"Academic Status (N = {total_lv})",
                tbl_lv,
            )

    # Add pie chart for each category (labels in legend to avoid overlapping), with sample sizes
    doc.add_heading('Charts', level=1)
    if not school_counts.empty:
        chart_title = f"Proportion of College/School (unique EIDs, N = {n_unique_eids})"
        doc.add_heading(chart_title, level=2)
        doc.add_picture(
            pie_chart_to_bytes(school_counts, chart_title),
            width=Inches(5.5),
        )
    for series, title in categories:
        counts = series.value_counts(dropna=False)
        # Match heading/sample size used in tables
        valid = series.dropna().astype(str).str.strip()
        valid_n = (valid != "").sum()
        chart_title = f"{title} (N = {valid_n})"
        doc.add_heading(chart_title, level=2)
        doc.add_picture(
            pie_chart_to_bytes(counts, chart_title),
            width=Inches(5.5),
        )

    if level_series is not None:
        level_counts = ordered_level_counts(level_series)
        if not level_counts.empty:
            chart_title = f"Academic Status (N = {n_unique_eids})"
            doc.add_heading(chart_title, level=2)
            doc.add_picture(
                pie_chart_to_bytes(level_counts, chart_title),
                width=Inches(5.5),
            )

    # Comparison to international enrollment (if reference file provided or found)
    enrollment_path = resolve_enrollment_path(csv_dir, script_dir, enrollment_reference_path)
    if enrollment_path and not school_counts.empty:
        enrollment_counts, total_enrollment = load_enrollment_by_school(enrollment_path, code_to_school)
        if total_enrollment > 0 and not enrollment_counts.empty:
            total_participants = n_unique_eids  # unique EIDs from event file (denominator for participation %)
            comparison_df = build_representation_comparison(
                school_counts, total_participants, enrollment_counts, total_enrollment
            )
            doc.add_heading("Comparison to International Enrollment (by School)", level=1)
            doc.add_paragraph(f"Enrollment reference: {enrollment_path}")
            comparison_table_to_doc(doc, comparison_df)
            doc.add_heading("Enrollment % vs Participation %", level=2)
            doc.add_picture(side_by_side_bar_chart_bytes(comparison_df), width=Inches(5.5))
            doc.add_heading("Representation Ratio & Over/Under by School", level=2)
            doc.add_picture(representation_ratio_chart_bytes(comparison_df, title="Representation Ratio & Over/Under by School"), width=Inches(5.5))
        else:
            print("Enrollment file found but no enrollment counts by school; skipping comparison section.")
    elif not school_counts.empty:
        print("No enrollment reference file found; comparison section omitted. Place All_International_Students_Enrolled.csv in the CSV dir or pass it as second argument.")

    country_enrollment_path = resolve_country_enrollment_path(csv_dir, script_dir)
    if country_col and country_counts is not None and not country_counts.empty:
        if country_enrollment_path:
            enc_c, tot_c, cont_map, name_map = load_enrollment_by_country(country_enrollment_path)
            if tot_c > 0 and not enc_c.empty:
                country_comp_df = build_country_representation_comparison(
                    country_counts,
                    n_unique_eids,
                    enc_c,
                    tot_c,
                    enrollment_display_by_key=name_map,
                    continent_by_key=cont_map,
                )
                if not country_comp_df.empty:
                    doc.add_heading("Comparison to International Enrollment (by Country)", level=1)
                    doc.add_paragraph(f"Country enrollment reference: {country_enrollment_path}")
                    country_intro = (
                        "Enrollment % = (country enrollment / total international enrollment in reference file) × 100. "
                        "Participation % = (participant EIDs with this country / total unique event EIDs) × 100. "
                        "Ratio = Participation % / Enrollment % (≈1 proportional, >1 overrepresented, <1 underrepresented)."
                    )
                    comparison_table_to_doc(doc, country_comp_df, intro_paragraph=country_intro)
                    doc.add_heading("Enrollment % vs Participation % by Country", level=2)
                    doc.add_picture(
                        side_by_side_bar_chart_bytes(
                            country_comp_df,
                            title="Enrollment % vs Participation % by Country (top 16 by enrollment %)",
                            label_column="Country",
                        ),
                        width=Inches(5.5),
                    )
                    doc.add_heading("Representation Ratio & Over/Under by Country", level=2)
                    doc.add_picture(
                        representation_ratio_chart_bytes(
                            country_comp_df,
                            title="Representation Ratio & Over/Under by Country (16 closest to proportional, by ratio)",
                            label_column="Country",
                        ),
                        width=Inches(5.5),
                    )
            else:
                print(
                    "Country enrollment reference file has no usable counts; skipping country comparison section."
                )
        else:
            print(
                "No country enrollment reference file found; country comparison omitted. "
                f"Place {COUNTRY_ENROLLMENT_REFERENCE_FILENAME} in the CSV dir, script dir, or Downloads."
            )

    try:
        doc.save(out_docx)
        print(f"Report saved to: {out_docx}")
        return out_docx
    except PermissionError:
        from datetime import datetime
        alt_name = f"{base}_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        alt_path = os.path.join(out_dir, alt_name)
        doc.save(alt_path)
        print(f"Original file is open or locked; report saved to: {alt_path}")
        return alt_path


def _looks_like_enrollment_reference_path(path):
    """Guess whether a path is the international enrollment extract (vs another participant export)."""
    b = os.path.basename(os.fspath(path)).lower()
    if "enroll" in b:
        return True
    if "international" in b and "student" in b:
        return True
    return False


def main():
    if len(sys.argv) < 2:
        print("Usage: python Generate_Proportion.py <event_participants.csv> [more.csv ...] [enrollment_reference.csv]")
        print("  Combine multiple participant lists: list each CSV in order (see below).")
        print("  Enrollment file: pass as the last argument if its name suggests enrollment (e.g. contains 'enroll' or 'international'+'student'),")
        print("  or use: python Generate_Proportion.py file1.csv file2.csv -- path/to/enrollment.csv")
        print("  If enrollment is omitted, looks for All_International_Students_Enrolled.csv in the first CSV's dir, script dir, or Downloads.")
        print(f"  Country over/under table: include Country on participants; uses {COUNTRY_ENROLLMENT_REFERENCE_FILENAME} from the same search locations.")
        sys.exit(1)

    args = sys.argv[1:]
    enrollment_explicit = None
    if "--" in args:
        i = args.index("--")
        event_args = args[:i]
        rest = args[i + 1 :]
        if not event_args:
            print("Error: no event CSV path(s) before --.")
            sys.exit(1)
        if rest:
            enrollment_explicit = rest[0]
            if len(rest) > 1:
                print("Note: ignoring extra arguments after enrollment path:", " ".join(rest[1:]))
    elif len(args) == 1:
        event_args = args
    elif len(args) == 2:
        # Backward compatible: event + enrollment
        event_args = [args[0]]
        enrollment_explicit = args[1]
    elif _looks_like_enrollment_reference_path(args[-1]):
        event_args = args[:-1]
        enrollment_explicit = args[-1]
    else:
        event_args = args

    generate_report(event_args, enrollment_explicit)

if __name__ == "__main__":
    main()
